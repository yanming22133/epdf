import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pypdf import PdfReader
import pdfplumber

class PdfToWord:
    def convert(self, pdf_path, output_path, preserve_layout=False):
        if preserve_layout:
            return self._convert_preserve_layout(pdf_path, output_path)
        else:
            return self._convert_text_only(pdf_path, output_path)

    def _convert_text_only(self, pdf_path, output_path):
        doc = Document()

        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    for paragraph_text in text.split('\n\n'):
                        if paragraph_text.strip():
                            p = doc.add_paragraph()
                            p.add_run(paragraph_text.strip())

                if page_num < len(pdf.pages) - 1:
                    doc.add_page_break()

        doc.save(output_path)
        return {'success': True, 'output': output_path}

    def _convert_preserve_layout(self, pdf_path, output_path):
        doc = Document()

        reader = PdfReader(pdf_path)

        for page_num, page in enumerate(reader.pages):
            blocks = page.get_text('dict')['blocks']

            for block in blocks:
                if 'lines' in block:
                    for line in block['lines']:
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                        for span in line['spans']:
                            text = span['text']
                            font_size = span.get('size', 12)
                            font_name = span.get('font', 'Times New Roman')

                            run = paragraph.add_run(text)

                            try:
                                run.font.size = Pt(font_size * 0.75)
                            except:
                                pass

                            try:
                                if 'Bold' in font_name:
                                    run.font.bold = True
                                if 'Italic' in font_name:
                                    run.font.italic = True
                            except:
                                pass

            if page_num < len(reader.pages) - 1:
                doc.add_page_break()

        doc.save(output_path)
        return {'success': True, 'output': output_path}

    def convert_with_tables(self, pdf_path, output_path):
        doc = Document()

        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                tables = page.extract_tables()

                if text:
                    for paragraph_text in text.split('\n\n'):
                        if paragraph_text.strip():
                            doc.add_paragraph(paragraph_text.strip())

                if tables:
                    for table in tables:
                        if table:
                            rows = len(table)
                            cols = len(table[0]) if table else 0
                            if rows > 0 and cols > 0:
                                t = doc.add_table(rows=rows, cols=cols)
                                t.style = 'Table Grid'
                                for i, row in enumerate(table):
                                    for j, cell in enumerate(row):
                                        t.cell(i, j).text = str(cell) if cell else ''

                if page_num < len(pdf.pages) - 1:
                    doc.add_page_break()

        doc.save(output_path)
        return {'success': True, 'output': output_path}

    def extract_text(self, pdf_path, page_start=None, page_end=None):
        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)

            start = (page_start - 1) if page_start else 0
            end = page_end if page_end else total_pages

            text = ''
            for page_num in range(start, min(end, total_pages)):
                page = pdf.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += f'\n--- Page {page_num + 1} ---\n'
                    text += page_text

        return text
